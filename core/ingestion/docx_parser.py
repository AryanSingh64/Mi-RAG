import math
from pathlib import Path
from typing import Any, Optional
from docx import Document
from core.ingestion.base import BaseDocumentParser, ParsedDocument


class DocxDocumentParser(BaseDocumentParser):
    """
    Parser for Microsoft Word (.docx) files.
    Extracts structured paragraphs and table cells, with two-way page slicing.
    """

    WORDS_PER_PAGE = 450
    PARAS_PER_PAGE = 15

    def parse(
        self,
        file_path: Path,
        start_page: Optional[int] = None,
        end_page: Optional[int] = None,
        progress_callback: Optional[Any] = None
    ) -> ParsedDocument:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        doc = Document(str(file_path))
        extracted_sections = []

        # 1. Extract body paragraphs (headings and text)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                extracted_sections.append(text)

        # 2. Extract tables (tabular enterprise data)
        for table_idx, table in enumerate(doc.tables, start=1):
            table_lines = [f"\n--- Table {table_idx} ---"]
            for row in table.rows:
                row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                table_lines.append(" | ".join(row_cells))
            if len(table_lines) > 1:
                extracted_sections.append("\n".join(table_lines))

        total_words = sum(len(s.split()) for s in extracted_sections)
        total_sections = len(extracted_sections)
        total_doc_pages = max(1, max(math.ceil(total_words / self.WORDS_PER_PAGE), math.ceil(total_sections / self.PARAS_PER_PAGE)))

        if start_page is not None or end_page is not None:
            s_p = max(1, start_page) if start_page is not None else 1
            e_p = min(total_doc_pages, end_page) if end_page is not None else total_doc_pages
            if s_p > e_p:
                s_p, e_p = 1, total_doc_pages

            s_idx = max(0, math.floor(((s_p - 1) / total_doc_pages) * total_sections))
            e_idx = min(total_sections, math.ceil((e_p / total_doc_pages) * total_sections))
            sliced_sections = extracted_sections[s_idx:e_idx] if total_sections > 0 else []
            full_text = "\n\n".join(sliced_sections)
            active_pages = max(1, e_p - s_p + 1)
        else:
            s_p, e_p = 1, total_doc_pages
            full_text = "\n\n".join(extracted_sections)
            active_pages = total_doc_pages

        return ParsedDocument(
            filename=file_path.name,
            file_path=str(file_path.resolve()),
            file_type="docx",
            text_content=full_text,
            metadata={
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "char_count": len(full_text),
                "total_pages": active_pages,
                "total_doc_pages": total_doc_pages,
                "page_range": f"{s_p}-{e_p}",
            }
        )
